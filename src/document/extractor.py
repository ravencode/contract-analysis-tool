"""
Document Extractor Module
Extracts text from PDF, DOCX, and TXT files
"""

import os
import re
from typing import Dict, Optional, Tuple
from pathlib import Path
from dataclasses import dataclass
import hashlib


@dataclass
class ExtractedDocument:
    """Extracted document data."""
    filename: str
    file_type: str
    text: str
    page_count: int
    word_count: int
    char_count: int
    file_hash: str
    metadata: Dict
    extraction_success: bool
    error_message: Optional[str] = None


class DocumentExtractor:
    """
    Extracts text content from various document formats.
    Supports PDF, DOCX, DOC, and TXT files.
    """
    
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx', '.doc', '.txt']
        
    def extract(self, file_path: str = None, 
                file_content: bytes = None,
                filename: str = None) -> ExtractedDocument:
        """
        Extract text from a document.
        
        Args:
            file_path: Path to the file (if reading from disk)
            file_content: File content as bytes (if uploaded)
            filename: Original filename (required if using file_content)
            
        Returns:
            ExtractedDocument with extracted text and metadata
        """
        try:
            # Determine file type
            if file_path:
                filename = os.path.basename(file_path)
                with open(file_path, 'rb') as f:
                    file_content = f.read()
            
            if not filename:
                raise ValueError("Filename is required")
            
            ext = Path(filename).suffix.lower()
            
            if ext not in self.supported_extensions:
                raise ValueError(f"Unsupported file type: {ext}")
            
            # Calculate file hash for audit
            file_hash = hashlib.sha256(file_content).hexdigest()[:16]
            
            # Extract based on file type
            if ext == '.pdf':
                text, page_count, metadata = self._extract_pdf(file_content)
            elif ext in ['.docx', '.doc']:
                text, page_count, metadata = self._extract_docx(file_content)
            elif ext == '.txt':
                text, page_count, metadata = self._extract_txt(file_content)
            else:
                raise ValueError(f"Unsupported file type: {ext}")
            
            # Clean the extracted text
            text = self._clean_text(text)
            
            return ExtractedDocument(
                filename=filename,
                file_type=ext,
                text=text,
                page_count=page_count,
                word_count=len(text.split()),
                char_count=len(text),
                file_hash=file_hash,
                metadata=metadata,
                extraction_success=True
            )
            
        except Exception as e:
            return ExtractedDocument(
                filename=filename or "unknown",
                file_type=ext if 'ext' in locals() else "unknown",
                text="",
                page_count=0,
                word_count=0,
                char_count=0,
                file_hash="",
                metadata={},
                extraction_success=False,
                error_message=str(e)
            )
    
    def _extract_pdf(self, content: bytes) -> Tuple[str, int, Dict]:
        """
        Extract text from PDF content.
        
        Args:
            content: PDF file content as bytes
            
        Returns:
            Tuple of (text, page_count, metadata)
        """
        import io
        
        text_parts = []
        page_count = 0
        metadata = {}
        
        # Try pdfplumber first (better for complex layouts)
        try:
            import pdfplumber
            
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                page_count = len(pdf.pages)
                metadata = pdf.metadata or {}
                
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            if text_parts:
                return '\n\n'.join(text_parts), page_count, metadata
        except Exception:
            pass
        
        # Fallback to PyPDF2
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(io.BytesIO(content))
            page_count = len(reader.pages)
            
            if reader.metadata:
                metadata = {
                    'title': reader.metadata.get('/Title', ''),
                    'author': reader.metadata.get('/Author', ''),
                    'creator': reader.metadata.get('/Creator', ''),
                    'creation_date': str(reader.metadata.get('/CreationDate', ''))
                }
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            return '\n\n'.join(text_parts), page_count, metadata
        except Exception as e:
            raise Exception(f"Failed to extract PDF: {str(e)}")
    
    def _extract_docx(self, content: bytes) -> Tuple[str, int, Dict]:
        """
        Extract text from DOCX content.
        
        Args:
            content: DOCX file content as bytes
            
        Returns:
            Tuple of (text, page_count, metadata)
        """
        import io
        from docx import Document
        
        try:
            doc = Document(io.BytesIO(content))
            
            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            # Extract metadata
            metadata = {}
            if doc.core_properties:
                metadata = {
                    'title': doc.core_properties.title or '',
                    'author': doc.core_properties.author or '',
                    'created': str(doc.core_properties.created or ''),
                    'modified': str(doc.core_properties.modified or '')
                }
            
            # Estimate page count (rough estimate)
            text = '\n\n'.join(text_parts)
            page_count = max(1, len(text) // 3000)
            
            return text, page_count, metadata
            
        except Exception as e:
            raise Exception(f"Failed to extract DOCX: {str(e)}")
    
    def _extract_txt(self, content: bytes) -> Tuple[str, int, Dict]:
        """
        Extract text from TXT content.
        
        Args:
            content: TXT file content as bytes
            
        Returns:
            Tuple of (text, page_count, metadata)
        """
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
        
        for encoding in encodings:
            try:
                text = content.decode(encoding)
                page_count = max(1, len(text) // 3000)
                return text, page_count, {'encoding': encoding}
            except UnicodeDecodeError:
                continue
        
        raise Exception("Failed to decode text file with supported encodings")
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove page numbers and headers/footers patterns
        text = re.sub(r'Page\s+\d+\s+of\s+\d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'-\s*\d+\s*-', '', text)
        
        # Normalize line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        
        # Normalize quotes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        
        return text.strip()
    
    def get_file_info(self, file_path: str = None,
                      file_content: bytes = None,
                      filename: str = None) -> Dict:
        """
        Get basic file information without full extraction.
        
        Args:
            file_path: Path to the file
            file_content: File content as bytes
            filename: Original filename
            
        Returns:
            Dictionary with file information
        """
        if file_path:
            filename = os.path.basename(file_path)
            file_size = os.path.getsize(file_path)
            with open(file_path, 'rb') as f:
                file_content = f.read()
        else:
            file_size = len(file_content) if file_content else 0
        
        ext = Path(filename).suffix.lower() if filename else ''
        file_hash = hashlib.sha256(file_content).hexdigest()[:16] if file_content else ''
        
        return {
            'filename': filename,
            'extension': ext,
            'size_bytes': file_size,
            'size_mb': round(file_size / (1024 * 1024), 2),
            'file_hash': file_hash,
            'is_supported': ext in self.supported_extensions
        }
    
    def validate_file(self, file_content: bytes, filename: str) -> Tuple[bool, str]:
        """
        Validate if a file can be processed.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (is_valid, message)
        """
        ext = Path(filename).suffix.lower()
        
        if ext not in self.supported_extensions:
            return False, f"Unsupported file type: {ext}. Supported: {', '.join(self.supported_extensions)}"
        
        # Check file size (max 10 MB)
        max_size = 10 * 1024 * 1024
        if len(file_content) > max_size:
            return False, f"File too large. Maximum size: 10 MB"
        
        # Check if file is empty
        if len(file_content) == 0:
            return False, "File is empty"
        
        # Basic validation for PDF
        if ext == '.pdf':
            if not file_content.startswith(b'%PDF'):
                return False, "Invalid PDF file"
        
        # Basic validation for DOCX
        if ext == '.docx':
            if not file_content.startswith(b'PK'):
                return False, "Invalid DOCX file"
        
        return True, "File is valid"
